import io
import re
import zipfile

import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st


try:
    import google.generativeai as genai
except ImportError:
    genai = None

try:
    import fitz
except ImportError:
    fitz = None

try:
    import faiss

except ImportError:
    faiss = None
    

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

try:
    import bibtexparser
except ImportError:
    bibtexparser = None

from jinja2 import Template
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

APP_TITLE = "AI Research Paper Studio"
GEMINI_MODEL = "gemini-2.5-flash-lite"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
TOP_K = 5
EMBED_MODEL_NAME = "all-MiniLM-L6-v2"

CITATION_STYLES = ["APA", "IEEE", "MLA", "Chicago"]
JOURNAL_TEMPLATES = ["IEEE", "Springer", "ACM", "Elsevier", "Generic"]
METHODOLOGY_TYPES = ["Experimental", "Survey", "Case Study", "Literature Review", "Comparative Study"]
CHART_TYPES = ["Bar Chart", "Line Chart", "Scatter Plot"]

SECTION_ORDER = [
    ("planner", "Research Planner"),
    ("literature_review", "Literature Review"),
    ("methodology", "Methodology"),
    ("results", "Results Analysis"),
    ("discussion", "Discussion"),
    ("conclusion", "Conclusion"),
    ("abstract", "Abstract"),
]


def init_state():
    defaults = {
        "step": 1,
        "api_key": "",
        "citation_style": "APA",
        "journal_template": "IEEE",
        "word_count_target": 4000,
        "project": {},
        "paper": {
            "title": "", "authors": [], "abstract": "", "problem_statement": "",
            "outline": "", "literature_review": "", "methodology": "",
            "results": "", "discussion": "", "conclusion": "", "references": [],
        },
        "ref_chunks": [],
        "ref_embeddings": None,
        "faiss_index": None,
        "ref_sources": [],
        "csv_df": None,
        "chart_fig": None,
        "review_feedback": "",
        "section_status": {k: False for k, _ in SECTION_ORDER},
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val


def get_paper():
    return st.session_state.paper


@st.cache_resource(show_spinner=False)
def load_embedding_model():
    if SentenceTransformer is None:
        return None
    return SentenceTransformer(EMBED_MODEL_NAME)


def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    chunks = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == n:
            break
        start = end - overlap
    return chunks


def extract_pdf_text(file_bytes):
    if fitz is None:
        return ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    except Exception:
        return ""


def build_faiss_index(chunks):
    model = load_embedding_model()
    if model is None or faiss is None or not chunks:
        return None, None
    embeddings = model.encode(chunks, show_progress_bar=False)
    embeddings = np.array(embeddings).astype("float32")
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(embeddings)
    return index, embeddings


def retrieve_chunks(query, k=TOP_K):
    index = st.session_state.faiss_index
    chunks = st.session_state.ref_chunks
    model = load_embedding_model()
    if index is None or model is None or not chunks:
        return []
    query_vec = np.array(model.encode([query])).astype("float32")
    k = min(k, len(chunks))
    distances, indices = index.search(query_vec, k)
    return [chunks[i] for i in indices[0] if i < len(chunks)]


def configure_gemini():
    if genai is None or not st.session_state.api_key:
        return False
    try:
        genai.configure(api_key=st.session_state.api_key)
        return True
    except Exception:
        return False


def call_gemini(prompt, max_tokens=1024, temperature=0.3):
    if not configure_gemini():
        return None, "Missing or invalid Gemini API key."
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": temperature,
                "max_output_tokens": max_tokens,
            },
        )
        return response.text.strip(), None
    except Exception as e:
        return None, f"Gemini request failed: {e}"


def project_context_block():
    p = st.session_state.project
    return (
        f"Title: {p.get('title', '')}\n"
        f"Domain: {p.get('domain', '')}\n"
        f"Keywords: {p.get('keywords', '')}\n"
        f"Objective: {p.get('objective', '')}\n"
        f"Research Questions: {p.get('questions', '')}\n"
        f"Methodology Type: {p.get('methodology_type', '')}"
    )


def generate_planner():
    prompt = (
        "You are an academic research planning assistant. Based on the project "
        "details below, produce concise sections: Problem Statement, Objectives, "
        "Research Scope, and a Paper Outline (list of sections with one-line "
        "descriptions). Be specific and avoid generic filler.\n\n"
        f"{project_context_block()}"
    )
    return call_gemini(prompt, max_tokens=900)


def generate_literature_review():
    query = f"{st.session_state.project.get('title', '')} {st.session_state.project.get('keywords', '')}"
    chunks = retrieve_chunks(query, TOP_K)
    if not chunks:
        context = "No reference material was uploaded or retrieved."
    else:
        context = "\n---\n".join(chunks)
    outline = get_paper().get("outline", "")
    prompt = (
        "You are writing a Literature Review section for a research paper. "
        "Use ONLY the reference excerpts provided below. Do not invent studies, "
        "authors, or findings. Compare and contrast findings across sources, "
        "and explicitly identify gaps in the literature. If the excerpts are "
        "insufficient to support a claim, state that the evidence is limited "
        "rather than fabricating content.\n\n"
        f"Outline context:\n{outline}\n\n"
        f"Reference excerpts:\n{context}"
    )
    return call_gemini(prompt, max_tokens=1200)


def generate_methodology():
    p = st.session_state.project
    prompt = (
        "Write a Methodology section for an academic research paper. Use the "
        "research questions and methodology type below to produce a realistic, "
        "appropriately detailed methodology (design, participants/data sources, "
        "instruments, procedure, analysis plan). Do not invent specific datasets, "
        "sample sizes, or experimental results that have not been provided.\n\n"
        f"Research Questions: {p.get('questions', '')}\n"
        f"Methodology Type: {p.get('methodology_type', '')}"
    )
    return call_gemini(prompt, max_tokens=900)


def generate_results(csv_summary=None):
    if csv_summary:
        prompt = (
            "Write a Results Analysis section for a research paper based ONLY "
            "on the dataset summary statistics below. Describe patterns, "
            "trends, and notable values factually. Do not invent any numbers "
            "not present in the summary.\n\n"
            f"Dataset summary:\n{csv_summary}"
        )
    else:
        prompt = (
            "Write an 'Expected Results' section for a research paper, framed "
            "as anticipated outcomes based on the methodology below, since no "
            "experimental data has been collected yet. Clearly mark this as "
            "expected/anticipated, not actual findings. Do not fabricate "
            "specific numeric results.\n\n"
            f"Methodology:\n{get_paper().get('methodology', '')}"
        )
    return call_gemini(prompt, max_tokens=900)


def generate_discussion():
    paper = get_paper()
    prompt = (
        "Write a Discussion section for a research paper. Base it strictly on "
        "the Literature Review, Methodology, and Results provided below. "
        "Discuss implications, limitations, and how the results relate to "
        "prior work. Do not introduce any new findings or data.\n\n"
        f"Literature Review:\n{paper.get('literature_review', '')}\n\n"
        f"Methodology:\n{paper.get('methodology', '')}\n\n"
        f"Results:\n{paper.get('results', '')}"
    )
    return call_gemini(prompt, max_tokens=900)


def generate_conclusion():
    paper = get_paper()
    prompt = (
        "Write a Conclusion section for a research paper. Summarize the "
        "problem, methodology, key results, and propose future work. Base it "
        "strictly on the content below; do not introduce new claims.\n\n"
        f"Problem Statement:\n{paper.get('problem_statement', '')}\n\n"
        f"Methodology:\n{paper.get('methodology', '')}\n\n"
        f"Results:\n{paper.get('results', '')}\n\n"
        f"Discussion:\n{paper.get('discussion', '')}"
    )
    return call_gemini(prompt, max_tokens=700)


def generate_abstract():
    paper = get_paper()
    prompt = (
        "Write a concise academic Abstract (150-250 words) summarizing the "
        "entire paper below: problem, methodology, key results, and "
        "conclusion. Do not introduce new information.\n\n"
        f"Problem Statement:\n{paper.get('problem_statement', '')}\n\n"
        f"Methodology:\n{paper.get('methodology', '')}\n\n"
        f"Results:\n{paper.get('results', '')}\n\n"
        f"Conclusion:\n{paper.get('conclusion', '')}"
    )
    return call_gemini(prompt, max_tokens=400)


def review_paper():
    paper = get_paper()
    full_text = "\n\n".join(
        f"{name}:\n{paper.get(key, '')}" for key, name in
        [("problem_statement", "Problem Statement"), ("literature_review", "Literature Review"),
         ("methodology", "Methodology"), ("results", "Results"),
         ("discussion", "Discussion"), ("conclusion", "Conclusion"), ("abstract", "Abstract")]
    )
    prompt = (
        "Review the academic paper draft below for: terminology consistency, "
        "repeated content across sections, missing sections, and citation "
        "inconsistencies. Return a short bulleted list of concrete improvement "
        "suggestions. Do not rewrite the paper.\n\n"
        f"{full_text}"
    )
    return call_gemini(prompt, max_tokens=600)


def parse_planner_output(text):
    sections = {"problem_statement": "", "outline": ""}
    problem_match = re.search(r"problem statement[:\n]*(.*?)(?=objectives|research scope|paper outline|$)", text, re.IGNORECASE | re.DOTALL)
    outline_match = re.search(r"paper outline[:\n]*(.*)", text, re.IGNORECASE | re.DOTALL)
    if problem_match:
        sections["problem_statement"] = problem_match.group(1).strip()
    else:
        sections["problem_statement"] = text[:400].strip()
    if outline_match:
        sections["outline"] = outline_match.group(1).strip()
    sections["full_text"] = text
    return sections


def summarize_csv(df):
    lines = [f"Rows: {len(df)}, Columns: {len(df.columns)}"]
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    for col in numeric_cols:
        desc = df[col].describe()
        lines.append(
            f"{col}: mean={desc['mean']:.2f}, std={desc['std']:.2f}, "
            f"min={desc['min']:.2f}, max={desc['max']:.2f}"
        )
    cat_cols = df.select_dtypes(exclude=[np.number]).columns
    for col in cat_cols[:5]:
        top_vals = df[col].value_counts().head(3).to_dict()
        lines.append(f"{col} top values: {top_vals}")
    return "\n".join(lines)


def build_chart(df, chart_type, x_col, y_col):
    try:
        if chart_type == "Bar Chart":
            fig = px.bar(df, x=x_col, y=y_col, template="simple_white")
        elif chart_type == "Line Chart":
            fig = px.line(df, x=x_col, y=y_col, template="simple_white")
        else:
            fig = px.scatter(df, x=x_col, y=y_col, template="simple_white")
        fig.update_layout(margin=dict(l=20, r=20, t=40, b=20))
        return fig
    except Exception:
        return None


def ref_field(ref, key, default=""):
    value = ref.get(key, default)
    if value is None or (isinstance(value, str) and not value.strip()):
        return default
    return value


def format_reference_apa(ref):
    authors = ref_field(ref, "authors")
    year = ref_field(ref, "year", "n.d.")
    title = ref_field(ref, "title")
    journal = ref_field(ref, "journal")
    return f"{authors} ({year}). {title}. {journal}."


def format_reference_ieee(ref, index):
    authors = ref_field(ref, "authors")
    title = ref_field(ref, "title")
    journal = ref_field(ref, "journal")
    year = ref_field(ref, "year", "n.d.")
    return f"[{index}] {authors}, \"{title},\" {journal}, {year}."


def format_reference_mla(ref):
    authors = ref_field(ref, "authors")
    title = ref_field(ref, "title")
    journal = ref_field(ref, "journal")
    year = ref_field(ref, "year", "n.d.")
    return f"{authors}. \"{title}.\" {journal}, {year}."


def format_reference_chicago(ref):
    authors = ref_field(ref, "authors")
    title = ref_field(ref, "title")
    journal = ref_field(ref, "journal")
    year = ref_field(ref, "year", "n.d.")
    return f"{authors}. \"{title}.\" {journal} ({year})."


def format_references(references, style):
    formatted = []
    for i, ref in enumerate(references, start=1):
        if not ref_field(ref, "title"):
            continue
        if style == "IEEE":
            formatted.append(format_reference_ieee(ref, i))
        elif style == "MLA":
            formatted.append(format_reference_mla(ref))
        elif style == "Chicago":
            formatted.append(format_reference_chicago(ref))
        else:
            formatted.append(format_reference_apa(ref))
    return formatted


LATEX_TEMPLATE = r"""
\documentclass[conference]{IEEEtran}
\usepackage{cite}
\usepackage{graphicx}
\begin{document}

\title{ {{ title }} }
\author{ {{ authors }} }
\maketitle

\begin{abstract}
{{ abstract }}
\end{abstract}

\section{Introduction}
{{ problem_statement }}

\section{Literature Review}
{{ literature_review }}

\section{Methodology}
{{ methodology }}

\section{Results}
{{ results }}

\section{Discussion}
{{ discussion }}

\section{Conclusion}
{{ conclusion }}

\bibliographystyle{IEEEtran}
\bibliography{references}

\end{document}
"""


def latex_escape(text):
    if not text:
        return ""
    replacements = {
        "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#",
        "_": r"\_", "{": r"\{", "}": r"\}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def build_bibtex(references):
    entries = []
    for i, ref in enumerate(references, start=1):
        if not ref_field(ref, "title"):
            continue
        key = f"ref{i}"
        entries.append(
            f"@article{{{key},\n"
            f"  author = {{{ref_field(ref, 'authors')}}},\n"
            f"  title = {{{ref_field(ref, 'title')}}},\n"
            f"  journal = {{{ref_field(ref, 'journal')}}},\n"
            f"  year = {{{ref_field(ref, 'year')}}}\n"
            f"}}\n"
        )
    return "\n".join(entries)


def build_latex_zip():
    paper = get_paper()
    template = Template(LATEX_TEMPLATE)
    tex_content = template.render(
        title=latex_escape(paper.get("title", "")),
        authors=latex_escape(", ".join(paper.get("authors", []))),
        abstract=latex_escape(paper.get("abstract", "")),
        problem_statement=latex_escape(paper.get("problem_statement", "")),
        literature_review=latex_escape(paper.get("literature_review", "")),
        methodology=latex_escape(paper.get("methodology", "")),
        results=latex_escape(paper.get("results", "")),
        discussion=latex_escape(paper.get("discussion", "")),
        conclusion=latex_escape(paper.get("conclusion", "")),
    )
    bib_content = build_bibtex(paper.get("references", []))

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("paper.tex", tex_content)
        zf.writestr("references.bib", bib_content)
        zf.writestr("figures/.gitkeep", "")
    buffer.seek(0)
    return buffer


def build_docx():
    paper = get_paper()
    doc = Document()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(paper.get("title", "Untitled Research Paper"))
    title_run.bold = True
    title_run.font.size = Pt(18)

    authors_p = doc.add_paragraph()
    authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_p.add_run(", ".join(paper.get("authors", [])))

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(paper.get("abstract", ""))

    sections = [
        ("Introduction", "problem_statement"),
        ("Literature Review", "literature_review"),
        ("Methodology", "methodology"),
        ("Results", "results"),
        ("Discussion", "discussion"),
        ("Conclusion", "conclusion"),
    ]
    for heading, key in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(paper.get(key, ""))

    doc.add_heading("References", level=1)
    style = st.session_state.citation_style
    for line in format_references(paper.get("references", []), style):
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def render_sidebar():
    with st.sidebar:
        st.header("Settings")
        st.session_state.api_key = st.text_input(
            "Gemini API Key", type="password", value=st.session_state.api_key,
            help="Your key is stored only in this browser session.",
        )
        st.session_state.citation_style = st.selectbox(
            "Citation Style", CITATION_STYLES,
            index=CITATION_STYLES.index(st.session_state.citation_style),
        )
        st.session_state.journal_template = st.selectbox(
            "Journal Template", JOURNAL_TEMPLATES,
            index=JOURNAL_TEMPLATES.index(st.session_state.journal_template),
        )
        st.session_state.word_count_target = st.number_input(
            "Word Count Target", min_value=1000, max_value=15000,
            value=st.session_state.word_count_target, step=500,
        )
        st.divider()
        st.caption("Export Settings")
        st.caption(f"LaTeX style: {st.session_state.journal_template}")
        st.caption(f"Citation style: {st.session_state.citation_style}")
        st.divider()
        completed = sum(st.session_state.section_status.values())
        st.progress(completed / len(SECTION_ORDER))
        st.caption(f"{completed}/{len(SECTION_ORDER)} sections generated")


def render_hero():
    st.title(APP_TITLE)
    st.caption("Guided, citation-safe research paper drafting for students.")
    cols = st.columns(4)
    cols[0].metric("Step", f"{st.session_state.step}/10")
    cols[1].metric("References", len(st.session_state.ref_sources))
    cols[2].metric("Sections Done", sum(st.session_state.section_status.values()))
    cols[3].metric("Word Target", st.session_state.word_count_target)
    st.divider()


def step_nav():
    cols = st.columns([1, 1, 6])
    with cols[0]:
        if st.button("Back", disabled=st.session_state.step <= 1, use_container_width=True):
            st.session_state.step -= 1
            st.rerun()
    with cols[1]:
        if st.button("Next", disabled=st.session_state.step >= 10, use_container_width=True):
            st.session_state.step += 1
            st.rerun()


def step_1_project_details():
    st.subheader("Step 1 - Project Details")
    p = st.session_state.project
    col1, col2 = st.columns(2)
    with col1:
        p["title"] = st.text_input("Paper Title", value=p.get("title", ""))
        p["domain"] = st.text_input("Research Domain", value=p.get("domain", ""))
        p["keywords"] = st.text_input("Keywords (comma separated)", value=p.get("keywords", ""))
        p["authors"] = st.text_input("Authors (comma separated)", value=p.get("authors", ""))
        p["affiliations"] = st.text_input("Affiliations", value=p.get("affiliations", ""))
    with col2:
        p["objective"] = st.text_area("Research Objective", value=p.get("objective", ""), height=100)
        p["questions"] = st.text_area("Research Questions", value=p.get("questions", ""), height=100)
        p["methodology_type"] = st.selectbox(
            "Methodology Type", METHODOLOGY_TYPES,
            index=METHODOLOGY_TYPES.index(p.get("methodology_type", METHODOLOGY_TYPES[0]))
            if p.get("methodology_type") in METHODOLOGY_TYPES else 0,
        )
    if st.button("Save Project Details", type="primary"):
        get_paper()["title"] = p.get("title", "")
        get_paper()["authors"] = [a.strip() for a in p.get("authors", "").split(",") if a.strip()]
        st.success("Project details saved.")


def step_2_upload_references():
    st.subheader("Step 2 - Upload References")
    files = st.file_uploader(
        "Upload PDF or TXT references", type=["pdf", "txt"], accept_multiple_files=True,
    )
    if files and st.button("Process References", type="primary"):
        all_chunks = []
        sources = []
        progress = st.progress(0.0)
        for i, f in enumerate(files):
            if f.name.lower().endswith(".pdf"):
                text = extract_pdf_text(f.read())
            else:
                text = f.read().decode("utf-8", errors="ignore")
            if not text.strip():
                st.warning(f"No extractable text found in {f.name}.")
                continue
            chunks = chunk_text(text)
            all_chunks.extend(chunks)
            sources.append(f.name)
            progress.progress((i + 1) / len(files))
        if all_chunks:
            with st.spinner("Building embedding index..."):
                index, embeddings = build_faiss_index(all_chunks)
            if index is None:
                st.error("Embedding index could not be built. Check that dependencies are installed.")
            else:
                st.session_state.ref_chunks = all_chunks
                st.session_state.faiss_index = index
                st.session_state.ref_embeddings = embeddings
                st.session_state.ref_sources = sources
                st.success(f"Indexed {len(all_chunks)} chunks from {len(sources)} file(s).")
        else:
            st.error("No usable text extracted from uploaded files.")

    if st.session_state.ref_sources:
        st.write("Indexed sources:")
        for s in st.session_state.ref_sources:
            st.caption(f"- {s}")

    st.divider()
    st.write("Manual Reference Entry")
    with st.form("manual_ref_form", clear_on_submit=True):
        cols = st.columns(2)
        title = cols[0].text_input("Title")
        authors = cols[1].text_input("Authors")
        cols2 = st.columns(3)
        year = cols2[0].text_input("Year")
        journal = cols2[1].text_input("Journal")
        doi = cols2[2].text_input("DOI")
        url = st.text_input("URL")
        if st.form_submit_button("Add Reference"):
            if title.strip():
                get_paper()["references"].append({
                    "title": title, "authors": authors, "year": year,
                    "journal": journal, "doi": doi, "url": url,
                })
                st.success("Reference added.")
            else:
                st.warning("Title is required to add a reference.")

    refs = get_paper().get("references", [])
    if refs:
        df = pd.DataFrame(refs)
        edited = st.data_editor(df, use_container_width=True, num_rows="dynamic")
        cleaned = edited.fillna("").to_dict("records")
        get_paper()["references"] = cleaned


def generation_block(label, session_key, generate_fn, help_text=""):
    st.write(f"**{label}**")
    if help_text:
        st.caption(help_text)
    paper = get_paper()
    existing = paper.get(session_key, "")
    button_label = f"Regenerate {label}" if existing else f"Generate {label}"
    if st.button(button_label, key=f"gen_{session_key}"):
        if not st.session_state.api_key:
            st.error("Please enter your Gemini API key in the sidebar first.")
        else:
            with st.spinner(f"Generating {label}..."):
                result, error = generate_fn()
            if error:
                st.error(error)
            else:
                paper[session_key] = result
                st.session_state.section_status[session_key] = True
                st.success(f"{label} generated.")
                st.rerun()
    if paper.get(session_key, ""):
        edited = st.text_area(
            f"Edit {label}", value=paper[session_key], height=250, key=f"edit_{session_key}",
        )
        paper[session_key] = edited
    elif not existing:
        st.info(f"{label} has not been generated yet.")


def step_3_outline():
    st.subheader("Step 3 - Generate Outline")

    def fn():
        text, error = generate_planner()
        if error:
            return None, error
        parsed = parse_planner_output(text)
        paper = get_paper()
        paper["problem_statement"] = parsed["problem_statement"]
        paper["outline"] = parsed["outline"] or parsed["full_text"]
        return parsed["full_text"], None

    generation_block(
        "Research Planner", "planner_output", fn,
        "Produces problem statement, objectives, scope, and outline from your project details.",
    )
    paper = get_paper()
    if paper.get("problem_statement"):
        with st.expander("Problem Statement", expanded=False):
            paper["problem_statement"] = st.text_area(
                "Edit Problem Statement", value=paper["problem_statement"], height=150,
            )
    if paper.get("outline"):
        with st.expander("Paper Outline", expanded=False):
            paper["outline"] = st.text_area("Edit Outline", value=paper["outline"], height=200)


def step_4_literature_review():
    st.subheader("Step 4 - Generate Literature Review")
    if not st.session_state.ref_chunks:
        st.warning("No references indexed yet. Upload references in Step 2 for grounded results.")
    generation_block(
        "Literature Review", "literature_review", generate_literature_review,
        "Uses only retrieved excerpts from your uploaded references.",
    )


def step_5_methodology():
    st.subheader("Step 5 - Generate Methodology")
    generation_block(
        "Methodology", "methodology", generate_methodology,
        "Based on your research questions and selected methodology type.",
    )


def step_6_results():
    st.subheader("Step 6 - Generate Results Analysis")
    csv_file = st.file_uploader("Upload CSV data (optional)", type=["csv"], key="csv_uploader")
    if csv_file is not None:
        try:
            df = pd.read_csv(csv_file)
            st.session_state.csv_df = df
            st.dataframe(df.head(10), use_container_width=True)
        except Exception as e:
            st.error(f"Invalid CSV file: {e}")

    df = st.session_state.csv_df
    if df is not None:
        cols = st.columns(3)
        chart_type = cols[0].selectbox("Chart Type", CHART_TYPES)
        x_col = cols[1].selectbox("X Axis", df.columns)
        y_col = cols[2].selectbox("Y Axis", df.columns)
        if st.button("Generate Chart"):
            fig = build_chart(df, chart_type, x_col, y_col)
            st.session_state.chart_fig = fig
        if st.session_state.chart_fig is not None:
            st.plotly_chart(st.session_state.chart_fig, use_container_width=True)

    def fn():
        if df is not None:
            summary = summarize_csv(df)
            return generate_results(csv_summary=summary)
        return generate_results(csv_summary=None)

    generation_block(
        "Results Analysis", "results", fn,
        "Uses your uploaded CSV summary statistics, or generates an Expected Results section if no data is provided.",
    )


def step_7_discussion():
    st.subheader("Step 7 - Generate Discussion")
    generation_block(
        "Discussion", "discussion", generate_discussion,
        "Synthesizes the Literature Review, Methodology, and Results without introducing new findings.",
    )


def step_8_conclusion():
    st.subheader("Step 8 - Generate Conclusion")
    generation_block(
        "Conclusion", "conclusion", generate_conclusion,
        "Summarizes the problem, methodology, results, and proposes future work.",
    )


def step_9_abstract():
    st.subheader("Step 9 - Generate Abstract")
    paper = get_paper()
    required = ["problem_statement", "methodology", "results", "discussion", "conclusion"]
    missing = [r for r in required if not paper.get(r)]
    if missing:
        st.warning(f"Generate these sections first: {', '.join(missing)}.")
    else:
        generation_block(
            "Abstract", "abstract", generate_abstract,
            "Generated only after all other sections exist; summarizes the full paper.",
        )


def step_10_compile_export():
    st.subheader("Step 10 - Compile and Export")
    paper = get_paper()

    tabs = st.tabs(["Preview", "Quality Check", "Export"])

    with tabs[0]:
        st.write(f"### {paper.get('title', 'Untitled')}")
        st.caption(", ".join(paper.get("authors", [])))
        section_map = [
            ("Abstract", "abstract"), ("Introduction", "problem_statement"),
            ("Literature Review", "literature_review"), ("Methodology", "methodology"),
            ("Results", "results"), ("Discussion", "discussion"), ("Conclusion", "conclusion"),
        ]
        for heading, key in section_map:
            with st.expander(heading, expanded=False):
                content = paper.get(key, "")
                if content:
                    st.write(content)
                else:
                    st.caption("Not generated yet.")
        refs = paper.get("references", [])
        with st.expander(f"References ({st.session_state.citation_style})", expanded=False):
            formatted = format_references(refs, st.session_state.citation_style)
            if formatted:
                for line in formatted:
                    st.write(line)
            else:
                st.caption("No references added.")

    with tabs[1]:
        st.write("Run an automated check for consistency and completeness.")
        if st.button("Review Paper", type="primary"):
            if not st.session_state.api_key:
                st.error("Please enter your Gemini API key in the sidebar first.")
            else:
                with st.spinner("Reviewing paper..."):
                    feedback, error = review_paper()
                if error:
                    st.error(error)
                else:
                    st.session_state.review_feedback = feedback
        if st.session_state.review_feedback:
            st.write(st.session_state.review_feedback)

    with tabs[2]:
        col1, col2 = st.columns(2)
        with col1:
            st.write("LaTeX Package")
            st.caption(f"Template: {st.session_state.journal_template}")
            try:
                zip_buffer = build_latex_zip()
                st.download_button(
                    "Download LaTeX ZIP", data=zip_buffer, file_name="research_paper_latex.zip",
                    mime="application/zip",
                )
            except Exception as e:
                st.error(f"LaTeX export failed: {e}")
        with col2:
            st.write("Word Document")
            try:
                docx_buffer = build_docx()
                st.download_button(
                    "Download DOCX", data=docx_buffer, file_name="research_paper.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.error(f"DOCX export failed: {e}")


STEP_RENDERERS = {
    1: step_1_project_details,
    2: step_2_upload_references,
    3: step_3_outline,
    4: step_4_literature_review,
    5: step_5_methodology,
    6: step_6_results,
    7: step_7_discussion,
    8: step_8_conclusion,
    9: step_9_abstract,
    10: step_10_compile_export,
}


def main():
    st.set_page_config(page_title=APP_TITLE, page_icon="📄", layout="wide")
    init_state()
    render_sidebar()
    render_hero()
    renderer = STEP_RENDERERS.get(st.session_state.step, step_1_project_details)
    renderer()
    st.divider()
    step_nav()


if __name__ == "__main__":
    main()